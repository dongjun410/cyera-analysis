#!/usr/bin/env python
"""DataDNA 分类引擎 — 主入口

Full pipeline orchestration: Tier 0 → Tier 1 (Stage A + B) → Tier 2 → Tier 3 → Discovery.

Usage:
    python main.py --input ./docs/ --output ./output/ --config config.yaml
"""

from __future__ import annotations

import argparse
import json
import logging
import time
from pathlib import Path
from typing import Any

import yaml

from src.discovery.loop import DiscoveryLoop
from src.embeddings.bge_m3 import BgeM3Embedder
from src.llm.client import LLMConfig, MistralClient
from src.ner.deberta import DebertaNER
from src.tier0.engine import Tier0Engine
from src.tier1.incremental import IncrementalAssigner
from src.tier1.semantic import SemanticRefiner
from src.tier1.structural import StructuralClusterer
from src.tier2.classifier import Tier2Classifier
from src.tier2.matching import KnownTypeMatcher
from src.tier2.propagation import LabelPropagator
from src.tier3.quality_gate import QualityGate
from src.types import ClassificationResult, ClusterInfo, Document

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# Document loading
# ═══════════════════════════════════════════════════════════════════

def load_documents(input_dir: str) -> list[Document]:
    """Walk input directory, read supported files, return Document objects.

    Supported formats: .txt, .pdf (pymupdf), .docx (python-docx), .json.
    Falls back to reading as plain text if a parser library is unavailable.
    """
    documents: list[Document] = []
    input_path = Path(input_dir)

    for file_path in sorted(input_path.rglob("*")):
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()
        if suffix not in (".txt", ".pdf", ".docx", ".json"):
            continue

        doc_id = file_path.stem
        rel_path = file_path.relative_to(input_path)
        metadata: dict[str, Any] = {
            "file_path": str(file_path),
            "file_type": suffix,
            "file_size": file_path.stat().st_size,
            "path_depth": max(len(rel_path.parts) - 1, 0),
        }

        text = ""
        try:
            if suffix == ".pdf":
                text = _read_pdf(file_path, metadata)
            elif suffix == ".docx":
                text = _read_docx(file_path, metadata)
            elif suffix == ".json":
                text = _read_json(file_path)
            else:
                text = file_path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            logger.warning("Failed to read %s: %s", file_path, exc)
            continue

        if not text.strip():
            logger.debug("Skipping empty document: %s", doc_id)
            continue

        documents.append(Document(doc_id=doc_id, text=text, metadata=metadata))

    logger.info("Loaded %d documents from %s", len(documents), input_dir)
    return documents


def _read_pdf(file_path: Path, metadata: dict[str, Any]) -> str:
    """Extract text from PDF using pymupdf. Falls back to plain-text read."""
    try:
        import fitz  # pymupdf
    except ImportError:
        logger.warning(
            "pymupdf not available, reading PDF as plain text: %s", file_path
        )
        return file_path.read_text(encoding="utf-8", errors="replace")

    doc = fitz.open(str(file_path))
    try:
        pages = [page.get_text() for page in doc]  # type: ignore[union-attr]
        metadata["page_count"] = len(pages)
        return "\n".join(pages)
    finally:
        doc.close()


def _read_docx(file_path: Path, metadata: dict[str, Any]) -> str:
    """Extract text from DOCX using python-docx. Falls back to plain-text read."""
    try:
        from docx import Document as DocxDocument  # type: ignore[assignment]
    except ImportError:
        logger.warning(
            "python-docx not available, reading DOCX as plain text: %s", file_path
        )
        return file_path.read_text(encoding="utf-8", errors="replace")

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs]
    metadata["paragraph_count"] = len(paragraphs)
    return "\n".join(paragraphs)


def _read_json(file_path: Path) -> str:
    """Read JSON file content, returning a formatted text representation."""
    try:
        data = json.loads(file_path.read_text(encoding="utf-8", errors="replace"))
        if isinstance(data, str):
            return data
        return json.dumps(data, ensure_ascii=False, indent=2)
    except (json.JSONDecodeError, UnicodeDecodeError):
        # Not valid JSON — read as raw text
        return file_path.read_text(encoding="utf-8", errors="replace")


# ═══════════════════════════════════════════════════════════════════
# Component initialization
# ═══════════════════════════════════════════════════════════════════

def _init_components(config: dict[str, Any]) -> dict[str, Any]:
    """Initialize all pipeline components from config sections.

    Returns a dict keyed by component name. Each component init is wrapped
    in its own try/except so a single failing component does not crash
    initialization — it is logged and re-raised (since the pipeline
    cannot run without all components).
    """
    components: dict[str, Any] = {}

    # ── Tier 0 ──────────────────────────────────────────────────
    try:
        tier0_config = config.get("tier0", {})
        components["engine"] = Tier0Engine(tier0_config)
        logger.info("Tier0Engine initialized")
    except Exception as exc:
        logger.error("Failed to initialize Tier0Engine: %s", exc)
        raise

    # ── Tier 1 Stage A: Structural hashing ──────────────────────
    try:
        stage_a_config = config.get("tier1", {}).get("stage_a", {})
        feature_keys = stage_a_config.get("structural_features")
        components["structural"] = StructuralClusterer(feature_config=feature_keys)
        logger.info("StructuralClusterer initialized")
    except Exception as exc:
        logger.error("Failed to initialize StructuralClusterer: %s", exc)
        raise

    # ── Embeddings ──────────────────────────────────────────────
    try:
        emb = config.get("embedding", {})
        components["embedder"] = BgeM3Embedder(
            model_name=emb.get("model_name", "BAAI/bge-m3"),
            device=emb.get("device", "cuda"),
            batch_size=emb.get("batch_size", 32),
            max_length=emb.get("max_token_length", 8192),
        )
        logger.info("BgeM3Embedder initialized (dim=%d)", components["embedder"].dim)
    except Exception as exc:
        logger.error("Failed to initialize BgeM3Embedder: %s", exc)
        raise

    # ── Tier 1 Stage B: Semantic refinement ─────────────────────
    try:
        stage_b_cfg = config.get("tier1", {}).get("stage_b", {})
        components["refiner"] = SemanticRefiner(components["embedder"], stage_b_cfg)
        logger.info("SemanticRefiner initialized")
    except Exception as exc:
        logger.error("Failed to initialize SemanticRefiner: %s", exc)
        raise

    # ── DeBERTa NER ─────────────────────────────────────────────
    try:
        t2 = config.get("tier2", {})
        ner_model = t2.get("ner_model", "microsoft/deberta-v3-base")
        ner_device = t2.get("ner_device", "cuda")
        components["ner"] = DebertaNER(model_name=ner_model, device=ner_device)
        logger.info("DebertaNER initialized")
    except Exception as exc:
        logger.error("Failed to initialize DebertaNER: %s", exc)
        raise

    # ── LLM Tier 2 (4-bit, 吞吐优先) ────────────────────────────
    try:
        t2_llm = t2.get("llm", {})
        components["llm_tier2"] = MistralClient(LLMConfig(
            api_base=t2_llm.get("api_base", "http://localhost:11434/v1"),
            model=t2_llm.get("model", "mistral:7b"),
            quantization=t2_llm.get("quantization", "4bit"),
            temperature=t2_llm.get("temperature", 0.3),
        ))
        logger.info("MistralClient (Tier2, 4-bit) initialized")
    except Exception as exc:
        logger.error("Failed to initialize MistralClient Tier2: %s", exc)
        raise

    # ── Known type matcher ──────────────────────────────────────
    try:
        match_cfg = t2.get("known_type_matching", {})
        components["matcher"] = KnownTypeMatcher(known_types=[], config=match_cfg)
        logger.info("KnownTypeMatcher initialized (0 known types)")
    except Exception as exc:
        logger.error("Failed to initialize KnownTypeMatcher: %s", exc)
        raise

    # ── Label propagator ────────────────────────────────────────
    try:
        prop_cfg = t2.get("propagation", {})
        components["propagator"] = LabelPropagator(prop_cfg)
        logger.info("LabelPropagator initialized")
    except Exception as exc:
        logger.error("Failed to initialize LabelPropagator: %s", exc)
        raise

    # ── Tier 2 classifier ───────────────────────────────────────
    try:
        components["classifier"] = Tier2Classifier(
            matcher=components["matcher"],
            ner=components["ner"],
            llm=components["llm_tier2"],
            propagator=components["propagator"],
        )
        logger.info("Tier2Classifier initialized")
    except Exception as exc:
        logger.error("Failed to initialize Tier2Classifier: %s", exc)
        raise

    # ── LLM Tier 3 (INT8, 精度优先) ────────────────────────────
    try:
        t3 = config.get("tier3", {})
        t3_llm = t3.get("llm", {})
        components["llm_tier3"] = MistralClient(LLMConfig(
            api_base=t3_llm.get("api_base", "http://localhost:11434/v1"),
            model=t3_llm.get("model", "mistral:7b"),
            quantization=t3_llm.get("quantization", "int8"),
            temperature=t3_llm.get("temperature", 0.1),
        ))
        logger.info("MistralClient (Tier3, INT8) initialized")
    except Exception as exc:
        logger.error("Failed to initialize MistralClient Tier3: %s", exc)
        raise

    # ── Quality gate ────────────────────────────────────────────
    try:
        components["quality_gate"] = QualityGate(components["llm_tier3"], t3)
        logger.info("QualityGate initialized")
    except Exception as exc:
        logger.error("Failed to initialize QualityGate: %s", exc)
        raise

    # ── Discovery loop ──────────────────────────────────────────
    try:
        disc_cfg = config.get("discovery", {})
        components["discovery"] = DiscoveryLoop(
            structural=components["structural"],
            refiner=components["refiner"],
            embedder=components["embedder"],
            matcher=components["matcher"],
            config=disc_cfg,
        )
        logger.info("DiscoveryLoop initialized")
    except Exception as exc:
        logger.error("Failed to initialize DiscoveryLoop: %s", exc)
        raise

    # ── Incremental assigner ────────────────────────────────────
    try:
        inc_cfg = config.get("tier1", {}).get("incremental", {})
        components["incremental"] = IncrementalAssigner(
            structural=components["structural"],
            refiner=components["refiner"],
            embedder=components["embedder"],
            config=inc_cfg,
        )
        logger.info("IncrementalAssigner initialized")
    except Exception as exc:
        logger.error("Failed to initialize IncrementalAssigner: %s", exc)
        raise

    return components


# ═══════════════════════════════════════════════════════════════════
# Output helpers
# ═══════════════════════════════════════════════════════════════════

def _result_to_dict(result: ClassificationResult) -> dict[str, Any]:
    """Convert a ClassificationResult to a JSON-serializable dict."""
    return {
        "doc_id": result.doc_id,
        "label": result.label,
        "confidence": result.confidence,
        "method": result.method,
        "is_new_type": result.is_new_type,
        "needs_manual_review": result.needs_manual_review,
        "rationale": result.rationale,
    }


def _write_output(output_dir: Path, results: list[dict[str, Any]], stats: dict[str, Any]) -> None:
    """Write results.json and stats.json to the output directory."""
    output_dir.mkdir(parents=True, exist_ok=True)

    results_path = output_dir / "results.json"
    with open(results_path, "w", encoding="utf-8") as fh:
        json.dump(results, fh, ensure_ascii=False, indent=2)
    logger.info("Results written to %s (%d entries)", results_path, len(results))

    stats_path = output_dir / "stats.json"
    with open(stats_path, "w", encoding="utf-8") as fh:
        json.dump(stats, fh, ensure_ascii=False, indent=2)
    logger.info("Stats written to %s", stats_path)


def _write_empty_output(output_dir: Path, stats: dict[str, Any]) -> None:
    """Write empty results when no documents are found."""
    stats["doc_count"] = 0
    stats["total_time_s"] = 0
    _write_output(output_dir, [], stats)


# ═══════════════════════════════════════════════════════════════════
# Main pipeline
# ═══════════════════════════════════════════════════════════════════

def main() -> int:
    parser = argparse.ArgumentParser(description="DataDNA 分类引擎")
    parser.add_argument("--input", required=True, help="文档目录路径")
    parser.add_argument("--output", default="./output/", help="输出目录")
    parser.add_argument("--config", default="config.yaml", help="配置文件路径")
    args = parser.parse_args()

    output_dir = Path(args.output)

    # ── Load config ──────────────────────────────────────────────
    with open(args.config, "r", encoding="utf-8") as fh:
        config: dict[str, Any] = yaml.safe_load(fh)

    logger.info("DataDNA 启动: input=%s, output=%s, config=%s",
                args.input, args.output, args.config)

    stats: dict[str, Any] = {}
    overall_start = time.perf_counter()

    # ── Initialize components ────────────────────────────────────
    init_start = time.perf_counter()
    comp = _init_components(config)
    stats["init_time_s"] = round(time.perf_counter() - init_start, 3)
    logger.info("All components initialized (%.3fs)", stats["init_time_s"])

    # ── Load documents ───────────────────────────────────────────
    try:
        doc_start = time.perf_counter()
        documents = load_documents(args.input)
        stats["doc_load_time_s"] = round(time.perf_counter() - doc_start, 3)
        stats["doc_count"] = len(documents)
    except Exception as exc:
        logger.error("Failed to load documents: %s", exc)
        return 1

    if not documents:
        logger.warning("No documents found in %s", args.input)
        _write_empty_output(output_dir, stats)
        logger.info("DataDNA 完成 (empty input)")
        return 0

    # ── Cold start check ─────────────────────────────────────────
    matcher: KnownTypeMatcher = comp["matcher"]
    classifier: Tier2Classifier = comp["classifier"]

    if matcher.type_count() == 0:
        logger.info("Cold start detected (%d docs, 0 known types) — "
                    "running zero-shot LLM classification", len(documents))
        try:
            results = classifier.cold_start_classify(documents)
            stats["method"] = "cold_start"
            stats["cold_start_docs"] = len(results)
            stats["total_time_s"] = round(time.perf_counter() - overall_start, 3)

            results_json = [_result_to_dict(r) for r in results]
            _write_output(output_dir, results_json, stats)
            logger.info("DataDNA 完成 (cold start): %d results, %.3fs",
                        len(results), stats["total_time_s"])
            return 0
        except Exception as exc:
            logger.error("Cold start classification failed: %s", exc)
            return 1

    # ── Tier 0: PII Feature Extraction ───────────────────────────
    engine: Tier0Engine = comp["engine"]
    stats["tier0"] = {}

    try:
        t0_start = time.perf_counter()
        doc_tuples: list[tuple[str, str]] = [(d.doc_id, d.text) for d in documents]
        pii_vectors = engine.extract_batch(doc_tuples)
        for doc, pii_vec in zip(documents, pii_vectors):
            doc.pii_features = pii_vec
        t0_time = round(time.perf_counter() - t0_start, 3)
        pii_detected = sum(1 for v in pii_vectors if v.pii_features)
        stats["tier0"] = {
            "time_s": t0_time,
            "docs_processed": len(documents),
            "docs_with_pii": pii_detected,
        }
        logger.info("Tier 0 complete: %d/%d docs with PII (%.3fs)",
                    pii_detected, len(documents), t0_time)
    except Exception as exc:
        logger.error("Tier 0 failed: %s — continuing without PII features", exc)
        stats["tier0"]["error"] = str(exc)

    # ── Tier 1 Stage A: Structural Clustering ────────────────────
    structural: StructuralClusterer = comp["structural"]
    refiner: SemanticRefiner = comp["refiner"]
    embedder: BgeM3Embedder = comp["embedder"]
    stage_b_config: dict[str, Any] = config.get("tier1", {}).get("stage_b", {})
    sem_split_threshold: int = stage_b_config.get("sem_split_threshold", 50)

    stats["tier1"] = {}
    buckets: dict[str, list[str]] = {}

    try:
        t1a_start = time.perf_counter()
        buckets = structural.cluster(documents)
        t1a_time = round(time.perf_counter() - t1a_start, 3)
        stats["tier1"]["stage_a_time_s"] = t1a_time
        stats["tier1"]["stage_a_buckets"] = len(buckets)
        logger.info("Tier 1 Stage A complete: %d structural buckets (%.3fs)",
                    len(buckets), t1a_time)
    except Exception as exc:
        logger.error("Tier 1 Stage A failed: %s — treating all docs as one bucket", exc)
        stats["tier1"]["stage_a_error"] = str(exc)
        # Fallback: single bucket with all documents
        fallback_id = "fallback_bucket"
        buckets = {fallback_id: [d.doc_id for d in documents]}

    # ── Tier 1 Stage B: Semantic Refinement ──────────────────────
    doc_lookup: dict[str, Document] = {d.doc_id: d for d in documents}
    all_clusters: list[ClusterInfo] = []

    try:
        t1b_start = time.perf_counter()
        for bucket_id, doc_ids in buckets.items():
            bucket_docs = [doc_lookup[did] for did in doc_ids if did in doc_lookup]
            if not bucket_docs:
                continue

            if len(bucket_docs) >= sem_split_threshold:
                # Large bucket — run semantic refinement
                sub_clusters = refiner.refine(bucket_id, bucket_docs)
                all_clusters.extend(sub_clusters)
            else:
                # Small bucket — treat as a single cluster
                cluster = ClusterInfo(
                    cluster_id=bucket_id,
                    doc_ids=sorted(bucket_docs, key=lambda d: d.doc_id),
                    structural_bucket=bucket_id,
                    cluster_radius=0.0,
                    representative_docs=[d.doc_id for d in bucket_docs[:3]],
                    tfidf_keywords=[],
                    pii_distribution={},
                    language_distribution={},
                )
                all_clusters.append(cluster)

        t1b_time = round(time.perf_counter() - t1b_start, 3)
        stats["tier1"]["stage_b_time_s"] = t1b_time
        stats["tier1"]["total_clusters"] = len(all_clusters)
        logger.info("Tier 1 Stage B complete: %d clusters total (%.3fs)",
                    len(all_clusters), t1b_time)
    except Exception as exc:
        logger.error("Tier 1 Stage B failed: %s", exc)
        stats["tier1"]["stage_b_error"] = str(exc)

    # ── Tier 2: Cluster Classification ───────────────────────────
    stats["tier2"] = {}
    classification_results: list[ClassificationResult] = []

    try:
        t2_start = time.perf_counter()
        classification_results = classifier.classify_clusters(all_clusters, documents)
        t2_time = round(time.perf_counter() - t2_start, 3)

        method_counts: dict[str, int] = {}
        for r in classification_results:
            method_counts[r.method] = method_counts.get(r.method, 0) + 1

        stats["tier2"] = {
            "time_s": t2_time,
            "results": len(classification_results),
            "method_breakdown": method_counts,
        }
        logger.info("Tier 2 complete: %d results, methods=%s (%.3fs)",
                    len(classification_results), method_counts, t2_time)
    except Exception as exc:
        logger.error("Tier 2 failed: %s", exc)
        stats["tier2"]["error"] = str(exc)
        # No fallback — classification results remain empty

    # Build lookups for Tier 3
    cluster_lookup: dict[str, ClusterInfo] = {c.cluster_id: c for c in all_clusters}
    doc_cluster_map: dict[str, str] = {}
    for c in all_clusters:
        for did in c.doc_ids:
            doc_cluster_map[did] = c.cluster_id

    # ── Tier 3: Quality Gate ─────────────────────────────────────
    quality_gate: QualityGate = comp["quality_gate"]
    stats["tier3"] = {}
    final_results: list[ClassificationResult] = list(classification_results)

    try:
        t3_start = time.perf_counter()
        tier3_triggered = 0

        for idx, result in enumerate(classification_results):
            doc = doc_lookup.get(result.doc_id)
            cluster_id = doc_cluster_map.get(result.doc_id)
            cluster = cluster_lookup.get(cluster_id or "")

            if doc is None or cluster is None:
                continue

            # Prepare trigger inputs
            ner_results_list: list[Any] = []
            tier0_features: dict[str, int] = {}
            if doc.pii_features is not None:
                tier0_features = doc.pii_features.pii_type_distribution

            if quality_gate.should_trigger(
                doc, cluster, result, ner_results_list, tier0_features,
            ):
                verified = quality_gate.verify(doc, cluster, result)
                final_results[idx] = verified
                tier3_triggered += 1

        t3_time = round(time.perf_counter() - t3_start, 3)
        total_classified = len(classification_results)
        trigger_rate = (
            round(tier3_triggered / total_classified, 4)
            if total_classified else 0.0
        )
        stats["tier3"] = {
            "time_s": t3_time,
            "triggered": tier3_triggered,
            "total": total_classified,
            "trigger_rate": trigger_rate,
        }
        logger.info("Tier 3 complete: %d/%d triggered (%.3fs, %.2f%%)",
                    tier3_triggered, total_classified, t3_time, trigger_rate * 100)
    except Exception as exc:
        logger.error("Tier 3 failed: %s — falling back to Tier 2 results", exc)
        stats["tier3"]["error"] = str(exc)
        final_results = classification_results

    # ── Discovery: Collect outliers ───────────────────────────────
    discovery: DiscoveryLoop = comp["discovery"]
    stats["discovery"] = {}

    try:
        disc_start = time.perf_counter()
        outlier_count = 0

        for result in final_results:
            doc = doc_lookup.get(result.doc_id)
            if doc is None:
                continue
            # Collect outliers: unknown classification, low confidence, or new type
            if result.method == "unknown" or result.confidence < 0.5 or result.is_new_type:
                discovery.collect_outlier(doc, f"low_confidence:{result.confidence:.2f}")
                outlier_count += 1

        stats["discovery"] = {
            "outliers_collected": outlier_count,
            "buffer_size": discovery.get_buffer_size(),
            "should_run": discovery.should_run(),
        }

        if discovery.should_run():
            new_types = discovery.run()
            stats["discovery"]["new_types_discovered"] = len(new_types)
            if new_types:
                logger.info("Discovery: registered %d new type(s)", len(new_types))

        stats["discovery"]["time_s"] = round(time.perf_counter() - disc_start, 3)
        logger.info("Discovery complete: %d outliers, buffer=%d, should_run=%s (%.3fs)",
                    outlier_count, discovery.get_buffer_size(),
                    discovery.should_run(), stats["discovery"]["time_s"])
    except Exception as exc:
        logger.error("Discovery failed: %s", exc)
        stats["discovery"]["error"] = str(exc)

    # ── Output ────────────────────────────────────────────────────
    stats["total_time_s"] = round(time.perf_counter() - overall_start, 3)

    try:
        results_json = [_result_to_dict(r) for r in final_results]
        _write_output(output_dir, results_json, stats)
    except Exception as exc:
        logger.error("Failed to write output: %s", exc)
        return 1

    logger.info("DataDNA 完成: %d results, %.3fs total",
                len(final_results), stats["total_time_s"])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
