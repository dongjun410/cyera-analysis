#!/usr/bin/env python
"""DataDNA 分类引擎 — 增量文档处理入口

单文档增量处理路径，用于初始聚类后到达的新文档分类。
依赖预存在的集群状态（state JSON），不运行完整的 Tier 1-3 管道。

Usage:
    python incremental.py --input ./new_docs/ --state ./output/state.json --output ./inc_output/ --config config.yaml
    python incremental.py --input single_doc.txt --state ./output/state.json --output ./inc_output/

工作流程:
  1. 加载集群状态 JSON → known_buckets (dict[bucket_id, list[ClusterInfo]])
  2. 加载新文档（单文件或目录）
  3. Tier 0: 提取 PII 特征
  4. IncrementalAssigner.assign() → 最近簇 或 离群
  5. 若分配成功 → 继承簇标签
  6. 若离群 → 标记为 "unclassified_outlier"，记录警告
  7. 检查离群比例是否触发重聚类
  8. 输出 results.json + stats.json
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
import time
from pathlib import Path
from typing import Any

import numpy as np
import yaml

from src.embeddings.bge_m3 import BgeM3Embedder
from src.llm.client import LLMConfig, MistralClient
from src.ner.deberta import DebertaNER
from src.tier0.engine import Tier0Engine
from src.tier1.incremental import IncrementalAssigner
from src.tier1.semantic import SemanticRefiner
from src.tier1.structural import StructuralClusterer
from src.tier2.classifier import Tier2Classifier
from src.tier2.matching import KnownTypeMatcher
from src.tier2.propagation import LabelPropagator
from src.tier3.quality_gate import QualityGate
from src.types import ClassificationResult, ClusterInfo, Document

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# Cluster state loading
# ═══════════════════════════════════════════════════════════════════

def load_cluster_state(state_path: str) -> dict[str, list[ClusterInfo]]:
    """从 JSON 加载集群状态，重建 ClusterInfo 对象。

    期望 JSON 结构::

        {
            "known_buckets": {
                "<bucket_id>": [
                    {
                        "cluster_id": "...",
                        "centroid_embedding": [0.1, 0.2, ...],
                        "cluster_radius": 0.5,
                        "label": "Medical Record",
                        "doc_ids": ["doc_01", ...],
                        "representative_docs": ["doc_01"],
                        "tfidf_keywords": ["patient", ...],
                        "pii_distribution": {"SSN": 5},
                        "language_distribution": {"en": 10},
                        "label_confidence": 0.9,
                        "needs_tier3": false
                    },
                    ...
                ],
                ...
            }
        }

    Returns:
        dict[bucket_id, list[ClusterInfo]] — 可直接传给 IncrementalAssigner.assign()
    """
    state_file = Path(state_path)
    if not state_file.exists():
        logger.error("集群状态文件不存在: %s", state_path)
        raise SystemExit(1)

    try:
        with open(state_file, "r", encoding="utf-8") as fh:
            state = json.load(fh)
    except (json.JSONDecodeError, OSError) as exc:
        logger.error("无法解析集群状态文件 %s: %s", state_path, exc)
        raise SystemExit(1)

    known_buckets_raw: dict[str, list[dict]] = state.get("known_buckets", {})
    known_buckets: dict[str, list[ClusterInfo]] = {}

    for bucket_id, cluster_list in known_buckets_raw.items():
        clusters: list[ClusterInfo] = []
        for entry in cluster_list:
            # 从 float 列表重建 centroid_embedding (np.ndarray)
            centroid_raw = entry.get("centroid_embedding")
            centroid: np.ndarray | None = None
            if centroid_raw is not None and len(centroid_raw) > 0:
                centroid = np.array(centroid_raw, dtype=np.float32)

            cluster = ClusterInfo(
                cluster_id=entry.get("cluster_id", ""),
                doc_ids=entry.get("doc_ids", []),
                structural_bucket=bucket_id,
                cluster_radius=float(entry.get("cluster_radius", 0.0)),
                representative_docs=entry.get("representative_docs", []),
                tfidf_keywords=entry.get("tfidf_keywords", []),
                pii_distribution=entry.get("pii_distribution", {}),
                language_distribution=entry.get("language_distribution", {}),
                centroid_embedding=centroid,
                label=entry.get("label"),
                label_confidence=float(entry.get("label_confidence", 0.0)),
                needs_tier3=bool(entry.get("needs_tier3", False)),
            )
            clusters.append(cluster)

        if clusters:
            known_buckets[bucket_id] = clusters

    total_clusters = sum(len(cl) for cl in known_buckets.values())
    logger.info(
        "已加载集群状态: %d 个结构桶, %d 个语义子簇",
        len(known_buckets), total_clusters,
    )
    return known_buckets


# ═══════════════════════════════════════════════════════════════════
# Document loading (single file + directory)
# ═══════════════════════════════════════════════════════════════════

SUPPORTED_SUFFIXES: set[str] = {".txt", ".pdf", ".docx", ".json"}


def load_single_document(file_path: str) -> Document | None:
    """加载单个文件为 Document 对象。

    Returns:
        Document 对象，或 None（不支持的格式/空文件/读取失败）。
    """
    path = Path(file_path)
    if not path.is_file():
        logger.error("输入文件不存在: %s", file_path)
        return None

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        logger.error("不支持的文件类型: %s (支持: %s)", suffix, SUPPORTED_SUFFIXES)
        return None

    doc_id = path.stem
    metadata: dict[str, Any] = {
        "file_path": str(path),
        "file_type": suffix,
        "file_size": path.stat().st_size,
        "path_depth": 0,
    }

    text = _read_file_content(path, suffix, metadata)
    if not text or not text.strip():
        logger.warning("空文档，跳过: %s", doc_id)
        return None

    return Document(doc_id=doc_id, text=text, metadata=metadata)


def load_documents_from_dir(input_dir: str) -> list[Document]:
    """递归遍历目录，加载所有支持的文档。

    按文件路径排序以保证确定性。
    """
    documents: list[Document] = []
    input_path = Path(input_dir)

    if not input_path.is_dir():
        logger.error("输入目录不存在: %s", input_dir)
        return documents

    for file_path in sorted(input_path.rglob("*")):
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            continue

        doc_id = file_path.stem
        rel_path = file_path.relative_to(input_path)
        metadata: dict[str, Any] = {
            "file_path": str(file_path),
            "file_type": suffix,
            "file_size": file_path.stat().st_size,
            "path_depth": max(len(rel_path.parts) - 1, 0),
        }

        try:
            text = _read_file_content(file_path, suffix, metadata)
        except Exception as exc:
            logger.warning("读取文件失败 %s: %s", file_path, exc)
            continue

        if not text or not text.strip():
            logger.debug("跳过空文档: %s", doc_id)
            continue

        documents.append(Document(doc_id=doc_id, text=text, metadata=metadata))

    logger.info("从 %s 加载了 %d 个文档", input_dir, len(documents))
    return documents


def _read_file_content(
    file_path: Path, suffix: str, metadata: dict[str, Any],
) -> str:
    """根据文件后缀调用对应的读取函数。"""
    if suffix == ".pdf":
        return _read_pdf(file_path, metadata)
    elif suffix == ".docx":
        return _read_docx(file_path, metadata)
    elif suffix == ".json":
        return _read_json(file_path)
    else:
        return file_path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(file_path: Path, metadata: dict[str, Any]) -> str:
    """使用 pymupdf 提取 PDF 文本。不可用时回退到纯文本读取。"""
    try:
        import fitz  # pymupdf
    except ImportError:
        logger.warning(
            "pymupdf 不可用，以纯文本方式读取 PDF: %s", file_path,
        )
        return file_path.read_text(encoding="utf-8", errors="replace")

    doc = fitz.open(str(file_path))
    try:
        pages = [page.get_text() for page in doc]  # type: ignore[union-attr]
        metadata["page_count"] = len(pages)
        return "\n".join(pages)
    finally:
        doc.close()


def _read_docx(file_path: Path, metadata: dict[str, Any]) -> str:
    """使用 python-docx 提取 DOCX 文本。不可用时回退到纯文本读取。"""
    try:
        from docx import Document as DocxDocument  # type: ignore[assignment]
    except ImportError:
        logger.warning(
            "python-docx 不可用，以纯文本方式读取 DOCX: %s", file_path,
        )
        return file_path.read_text(encoding="utf-8", errors="replace")

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs]
    metadata["paragraph_count"] = len(paragraphs)
    return "\n".join(paragraphs)


def _read_json(file_path: Path) -> str:
    """读取 JSON 文件内容，转为格式化的文本表示。"""
    try:
        data = json.loads(file_path.read_text(encoding="utf-8", errors="replace"))
        if isinstance(data, str):
            return data
        return json.dumps(data, ensure_ascii=False, indent=2)
    except (json.JSONDecodeError, UnicodeDecodeError):
        return file_path.read_text(encoding="utf-8", errors="replace")


# ═══════════════════════════════════════════════════════════════════
# Component initialization (增量路径最小集合)
# ═══════════════════════════════════════════════════════════════════

def _init_components(config: dict[str, Any]) -> dict[str, Any]:
    """初始化增量处理所需的组件集合。

    组件: Tier0Engine, StructuralClusterer, BgeM3Embedder,
           SemanticRefiner, IncrementalAssigner,
           DebertaNER, MistralClient (Tier2 + Tier3),
           KnownTypeMatcher, LabelPropagator,
           Tier2Classifier (outlier per-doc classification),
           QualityGate (outlier verification)
    """
    components: dict[str, Any] = {}

    # ── Tier 0 ──────────────────────────────────────────────────
    tier0_config = config.get("tier0", {})
    components["engine"] = Tier0Engine(tier0_config)
    logger.info("Tier0Engine 已初始化")

    # ── Tier 1 Stage A: 结构哈希 ────────────────────────────────
    stage_a_config = config.get("tier1", {}).get("stage_a", {})
    feature_keys = stage_a_config.get("structural_features")
    components["structural"] = StructuralClusterer(feature_config=feature_keys)
    logger.info("StructuralClusterer 已初始化")

    # ── BGE-M3 嵌入模型 ────────────────────────────────────────
    emb = config.get("embedding", {})
    components["embedder"] = BgeM3Embedder(
        model_name=emb.get("model_name", "BAAI/bge-m3"),
        device=emb.get("device", "cuda"),
        batch_size=emb.get("batch_size", 32),
        max_length=emb.get("max_token_length", 8192),
    )
    logger.info("BgeM3Embedder 已初始化 (dim=%d)", components["embedder"].dim)

    # ── SemanticRefiner (IncrementalAssigner 构造函数需要) ──────
    stage_b_cfg = config.get("tier1", {}).get("stage_b", {})
    components["refiner"] = SemanticRefiner(components["embedder"], stage_b_cfg)
    logger.info("SemanticRefiner 已初始化")

    # ── IncrementalAssigner ─────────────────────────────────────
    inc_cfg = config.get("tier1", {}).get("incremental", {})
    components["incremental"] = IncrementalAssigner(
        structural=components["structural"],
        refiner=components["refiner"],
        embedder=components["embedder"],
        config=inc_cfg,
    )
    logger.info("IncrementalAssigner 已初始化")

    # ── DeBERTa NER (用于离群文档分类) ──────────────────────────
    t2 = config.get("tier2", {})
    ner_model = t2.get("ner_model", "microsoft/deberta-v3-base")
    ner_device = t2.get("ner_device", "cuda")
    components["ner"] = DebertaNER(model_name=ner_model, device=ner_device)
    logger.info("DebertaNER 已初始化")

    # ── LLM Tier 2 (4-bit, 用于离群文档逐文档分类) ────────────
    t2_llm = t2.get("llm", {})
    components["llm_tier2"] = MistralClient(LLMConfig(
        api_base=t2_llm.get("api_base", "http://localhost:11434/v1"),
        model=t2_llm.get("model", "mistral:7b"),
        quantization=t2_llm.get("quantization", "4bit"),
        temperature=t2_llm.get("temperature", 0.3),
    ))
    logger.info("MistralClient (Tier2) 已初始化")

    # ── KnownTypeMatcher ────────────────────────────────────────
    match_cfg = t2.get("known_type_matching", {})
    components["matcher"] = KnownTypeMatcher(known_types=[], config=match_cfg)
    logger.info("KnownTypeMatcher 已初始化")

    # ── LabelPropagator ─────────────────────────────────────────
    prop_cfg = t2.get("propagation", {})
    components["propagator"] = LabelPropagator(prop_cfg)
    logger.info("LabelPropagator 已初始化")

    # ── Tier2Classifier (离群文档逐文档分类) ────────────────────
    components["classifier"] = Tier2Classifier(
        matcher=components["matcher"],
        ner=components["ner"],
        llm=components["llm_tier2"],
        propagator=components["propagator"],
    )
    logger.info("Tier2Classifier 已初始化")

    # ── LLM Tier 3 (INT8, 离群文档质量验证) ────────────────────
    t3 = config.get("tier3", {})
    t3_llm = t3.get("llm", {})
    components["llm_tier3"] = MistralClient(LLMConfig(
        api_base=t3_llm.get("api_base", "http://localhost:11434/v1"),
        model=t3_llm.get("model", "mistral:7b"),
        quantization=t3_llm.get("quantization", "int8"),
        temperature=t3_llm.get("temperature", 0.1),
    ))
    logger.info("MistralClient (Tier3) 已初始化")

    # ── QualityGate ─────────────────────────────────────────────
    components["quality_gate"] = QualityGate(components["llm_tier3"], t3)
    logger.info("QualityGate 已初始化")

    return components


# ═══════════════════════════════════════════════════════════════════
# Output helpers
# ═══════════════════════════════════════════════════════════════════

def _write_output(
    output_dir: Path,
    results: list[dict[str, Any]],
    stats: dict[str, Any],
) -> None:
    """写入 results.json 和 stats.json 到输出目录。"""
    output_dir.mkdir(parents=True, exist_ok=True)

    results_path = output_dir / "results.json"
    with open(results_path, "w", encoding="utf-8") as fh:
        json.dump(results, fh, ensure_ascii=False, indent=2)
    logger.info("结果已写入 %s (%d 条)", results_path, len(results))

    stats_path = output_dir / "stats.json"
    with open(stats_path, "w", encoding="utf-8") as fh:
        json.dump(stats, fh, ensure_ascii=False, indent=2)
    logger.info("统计已写入 %s", stats_path)


def _write_empty_output(output_dir: Path, stats: dict[str, Any]) -> None:
    """写入空结果（无文档或有错误时）。"""
    stats["doc_count"] = 0
    stats["assigned_count"] = 0
    stats["outlier_count"] = 0
    stats["total_time_s"] = 0
    _write_output(output_dir, [], stats)


# ═══════════════════════════════════════════════════════════════════
# Main pipeline
# ═══════════════════════════════════════════════════════════════════

def main() -> int:
    parser = argparse.ArgumentParser(
        description="DataDNA 增量文档处理 — 对初始聚类后的新文档进行分类",
    )
    parser.add_argument(
        "--input", required=True,
        help="单个文件路径 或 文档目录路径",
    )
    parser.add_argument(
        "--state", required=True,
        help="集群状态 JSON 文件路径 (包含 known_buckets)",
    )
    parser.add_argument(
        "--output", default="./output_incremental/",
        help="输出目录 (默认: ./output_incremental/)",
    )
    parser.add_argument(
        "--config", default="config.yaml",
        help="配置文件路径 (默认: config.yaml)",
    )
    args = parser.parse_args()

    output_dir = Path(args.output)

    # ── 加载配置 ──────────────────────────────────────────────────
    try:
        with open(args.config, "r", encoding="utf-8") as fh:
            config: dict[str, Any] = yaml.safe_load(fh)
    except FileNotFoundError:
        logger.error("配置文件不存在: %s", args.config)
        return 1
    except yaml.YAMLError as exc:
        logger.error("配置文件 YAML 解析失败: %s", exc)
        return 1

    logger.info(
        "DataDNA 增量处理启动: input=%s, state=%s, output=%s, config=%s",
        args.input, args.state, args.output, args.config,
    )

    stats: dict[str, Any] = {}
    overall_start = time.perf_counter()

    # ── 加载集群状态 ──────────────────────────────────────────────
    try:
        known_buckets = load_cluster_state(args.state)
    except SystemExit:
        return 1

    total_clusters = sum(len(cl) for cl in known_buckets.values())
    stats["known_buckets"] = len(known_buckets)
    stats["known_clusters"] = total_clusters

    if not known_buckets:
        logger.warning(
            "集群状态为空 (0 个已知桶) — 所有文档将被标记为 new_structure_candidate"
        )

    # ── 初始化组件 ────────────────────────────────────────────────
    init_start = time.perf_counter()
    try:
        comp = _init_components(config)
    except Exception as exc:
        logger.error("组件初始化失败: %s", exc)
        stats["init_error"] = str(exc)
        _write_empty_output(output_dir, stats)
        return 1
    stats["init_time_s"] = round(time.perf_counter() - init_start, 3)
    logger.info("所有组件已初始化 (%.3fs)", stats["init_time_s"])

    # ── 加载文档 ──────────────────────────────────────────────────
    input_path = Path(args.input)
    doc_start = time.perf_counter()

    if input_path.is_file():
        doc = load_single_document(args.input)
        documents = [doc] if doc is not None else []
        logger.info("单文件模式: %s", args.input)
    elif input_path.is_dir():
        documents = load_documents_from_dir(args.input)
    else:
        logger.error("输入路径不存在: %s", args.input)
        return 1

    stats["doc_load_time_s"] = round(time.perf_counter() - doc_start, 3)
    stats["doc_count"] = len(documents)

    if not documents:
        logger.warning("在 %s 中未找到文档", args.input)
        _write_empty_output(output_dir, stats)
        logger.info("DataDNA 增量处理完成 (空输入)")
        return 0

    # ── 逐文档处理 ────────────────────────────────────────────────
    engine: Tier0Engine = comp["engine"]
    assigner: IncrementalAssigner = comp["incremental"]
    classifier: Tier2Classifier = comp["classifier"]
    quality_gate: QualityGate = comp["quality_gate"]

    results: list[dict[str, Any]] = []
    assigned_count = 0
    outlier_count = 0
    error_count = 0

    # 构建 cluster_id → label 的快速查找表
    cluster_label_map: dict[str, str] = {}
    for bucket_clusters in known_buckets.values():
        for cluster in bucket_clusters:
            if cluster.label:
                cluster_label_map[cluster.cluster_id] = cluster.label

    t0_start = time.perf_counter()

    for doc in documents:
        # Step 3: Tier 0 — PII 特征提取
        try:
            pii_vec = engine.extract(doc.doc_id, doc.text)
            doc.pii_features = pii_vec
        except Exception as exc:
            logger.warning(
                "Tier 0 特征提取失败 doc=%s: %s — 继续处理", doc.doc_id, exc,
            )

        # Step 4: IncrementalAssigner.assign()
        try:
            assignment = assigner.assign(doc, known_buckets)
        except Exception as exc:
            logger.error(
                "文档分配失败 doc=%s: %s — 标记为未分类离群", doc.doc_id, exc,
            )
            error_count += 1
            outlier_count += 1
            results.append({
                "doc_id": doc.doc_id,
                "label": "unclassified_outlier",
                "confidence": 0.0,
                "method": "incremental_error",
                "is_new_type": False,
                "needs_manual_review": True,
                "rationale": f"Assignment error: {exc}",
            })
            continue

        if assignment.assigned_cluster_id is not None:
            # Step 5: 分配成功 → 继承簇标签
            assigned_count += 1
            label = cluster_label_map.get(
                assignment.assigned_cluster_id, "unknown",
            )
            results.append({
                "doc_id": doc.doc_id,
                "label": label,
                "confidence": 0.85,
                "method": "incremental_assignment",
                "is_new_type": False,
                "needs_manual_review": False,
                "rationale": (
                    f"Assigned to cluster {assignment.assigned_cluster_id}"
                ),
            })
        else:
            # Step 6: 离群 → 运行 Tier 2+3 逐文档分类
            outlier_count += 1
            logger.warning(
                "检测到离群文档: doc=%s, reason=%s, needs_recluster=%s",
                doc.doc_id, assignment.outlier_reason, assignment.needs_reclustering,
            )

            # Tier 2: Per-doc zero-shot LLM classification
            try:
                tier2_results = classifier.cold_start_classify([doc])
                if tier2_results:
                    cr = tier2_results[0]
                else:
                    cr = ClassificationResult(
                        doc_id=doc.doc_id,
                        label="unclassified_outlier",
                        confidence=0.0,
                        method="incremental_outlier",
                        is_new_type=True,
                        needs_manual_review=True,
                        rationale="Tier 2 returned no results",
                    )
            except Exception as exc:
                logger.error(
                    "Tier 2 离群分类失败 doc=%s: %s", doc.doc_id, exc,
                )
                cr = ClassificationResult(
                    doc_id=doc.doc_id,
                    label="unclassified_outlier",
                    confidence=0.0,
                    method="incremental_outlier",
                    is_new_type=True,
                    needs_manual_review=True,
                    rationale=f"Tier 2 error: {exc}",
                )

            # Tier 3: Quality gate verification for sensitive/low-confidence results
            try:
                if quality_gate.should_trigger(
                    doc,
                    ClusterInfo(
                        cluster_id="outlier-synthetic",
                        doc_ids=[doc.doc_id],
                        structural_bucket="outlier",
                        cluster_radius=0.0,
                        representative_docs=[doc.doc_id],
                        tfidf_keywords=[],
                        pii_distribution=(
                            doc.pii_features.pii_type_distribution
                            if doc.pii_features else {}
                        ),
                        language_distribution={},
                    ),
                    cr,
                ):
                    cr = quality_gate.verify(
                        doc,
                        ClusterInfo(
                            cluster_id="outlier-synthetic",
                            doc_ids=[doc.doc_id],
                            structural_bucket="outlier",
                            cluster_radius=0.0,
                            representative_docs=[doc.doc_id],
                            tfidf_keywords=[],
                            pii_distribution=(
                                doc.pii_features.pii_type_distribution
                                if doc.pii_features else {}
                            ),
                            language_distribution={},
                        ),
                        cr,
                    )
                    cr.method = "llm_tier3"
            except Exception as exc:
                logger.warning(
                    "Tier 3 验证失败 doc=%s: %s — 使用 Tier 2 结果", doc.doc_id, exc,
                )

            results.append({
                "doc_id": cr.doc_id,
                "label": cr.label,
                "confidence": cr.confidence,
                "method": cr.method,
                "is_new_type": cr.is_new_type,
                "needs_manual_review": cr.needs_manual_review,
                "rationale": cr.rationale,
            })

    stats["tier0"] = {
        "time_s": round(time.perf_counter() - t0_start, 3),
        "docs_processed": len(documents),
    }

    # Step 7: 检查离群比例是否触发重聚类
    recluster_warnings: list[dict[str, Any]] = []
    for bucket_id in sorted(known_buckets):
        if assigner.should_recluster(bucket_id):
            outlier_in_bucket = assigner.get_outlier_count(bucket_id)
            bucket_total = assigner.get_bucket_total(bucket_id)
            ratio = outlier_in_bucket / bucket_total if bucket_total > 0 else 0.0
            msg = (
                f"触发重聚类警告: bucket={bucket_id[:16]}... "
                f"离群={outlier_in_bucket}/{bucket_total} ({ratio:.1%})"
            )
            logger.warning(msg)
            recluster_warnings.append({
                "bucket_id": bucket_id,
                "outlier_count": outlier_in_bucket,
                "bucket_total": bucket_total,
                "outlier_ratio": round(ratio, 4),
            })

    if recluster_warnings:
        stats["recluster_warnings"] = recluster_warnings
        logger.warning(
            "共 %d 个桶触发重聚类警告", len(recluster_warnings),
        )

    # ── 输出 ──────────────────────────────────────────────────────
    stats["assigned_count"] = assigned_count
    stats["outlier_count"] = outlier_count
    stats["error_count"] = error_count
    stats["total_time_s"] = round(time.perf_counter() - overall_start, 3)

    try:
        _write_output(output_dir, results, stats)
    except Exception as exc:
        logger.error("输出写入失败: %s", exc)
        return 1

    logger.info(
        "DataDNA 增量处理完成: assigned=%d, outliers=%d, errors=%d, "
        "total=%d (%.3fs)",
        assigned_count, outlier_count, error_count,
        len(results), stats["total_time_s"],
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
